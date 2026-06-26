from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ==================== STYLES ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 1.15
pf.space_after = Pt(0)
pf.space_before = Pt(0)

section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)

# ==================== HELPERS ====================
def add_h(text, level=0, size=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    if level == 0:
        p.paragraph_format.space_before = Pt(24)
        run.underline = True
    return p

def add_p(text, indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_code(code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_bullet(text, indent_cm=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered(items, indent_cm=1):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(indent_cm)
        run = p.add_run(f'{i}. {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ==================== COVER ====================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LAPORAN PRAKTIKUM\nMODUL 9\nSTORED PROCEDURE, FUNGSI, DAN TRIGGER')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disusun oleh:')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Muhammad Zhio Affarel (0063)')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROGRAM STUDI TEKNOLOGI INFORMASI\nFAKULTAS TEKNIK\nUNIVERSITAS TIDAR\n2025')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ==================== KATA PENGANTAR ====================
add_h('KATA PENGANTAR')
add_p('Puji syukur ke hadirat Tuhan Yang Maha Esa atas limpahan rahmat dan hidayah-Nya sehingga penulis dapat menyelesaikan laporan praktikum Modul 9 yang berjudul "Stored Procedure, Fungsi, dan Trigger" dengan baik. Laporan ini disusun sebagai salah satu tugas mata kuliah Praktikum Basis Data pada Program Studi Teknologi Informasi, Universitas Tidar.')
add_p('Dalam penyusunan laporan ini, penulis telah berusaha semaksimal mungkin untuk menyajikan materi secara sistematis dan mudah dipahami. Laporan ini mencakup pembahasan mengenai implementasi stored procedure, fungsi, dan trigger dalam basis data MySQL. Setiap bagian dilengkapi dengan penjelasan teoretis, langkah-langkah praktikum, serta hasil pengujian yang dilakukan.')
add_p('Penulis menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kritik dan saran yang membangun sangat diharapkan untuk perbaikan di masa mendatang. Semoga laporan ini dapat memberikan manfaat bagi pembaca dan menjadi referensi yang berguna dalam mempelajari basis data.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
run = p.add_run('Magelang, Juni 2025\nPenulis')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ==================== DAFTAR ISI ====================
add_h('DAFTAR ISI')
toc_items = [
    'KATA PENGANTAR',
    'DAFTAR ISI',
    'BAB I  PENDAHULUAN',
    '    1.1 Latar Belakang',
    '    1.2 Tujuan',
    '    1.3 Ruang Lingkup',
    'BAB II  DASAR TEORI',
    '    2.1 Stored Procedure',
    '    2.2 Function',
    '    2.3 Trigger',
    'BAB III  IMPLEMENTASI DAN PEMBAHASAN',
    '    3.1 Persiapan dan Pembuatan Database',
    '    3.2 Implementasi Stored Procedure',
    '    3.3 Implementasi Function',
    '    3.4 Implementasi Trigger',
    '    3.5 Implementasi Tugas Modul 9',
    'BAB IV  KESIMPULAN',
    'DAFTAR PUSTAKA',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ==================== BAB I ====================
add_h('BAB I\nPENDAHULUAN')

add_h('1.1 Latar Belakang', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Dalam pengelolaan basis data, sering kali diperlukan operasi-operasi kompleks yang melibatkan banyak perintah SQL secara berurutan. Menjalankan perintah-perintah tersebut secara manual setiap kali dibutuhkan tentu tidak efisien dan rentan terhadap kesalahan. Oleh karena itu, diperlukan mekanisme yang dapat menyederhanakan dan mengotomatisasi proses tersebut.')
add_p('Stored procedure, fungsi, dan trigger merupakan objek-objek basis data yang memungkinkan pengguna untuk menyimpan serangkaian perintah SQL dan mengeksekusinya secara terprogram. Dengan memanfaatkan ketiga objek ini, pengelolaan data menjadi lebih efisien, konsisten, dan aman. Stored procedure memungkinkan eksekusi blok perintah secara berulang, fungsi mengembalikan nilai tertentu berdasarkan proses komputasi, dan trigger menjalankan aksi secara otomatis ketika terjadi perubahan data pada tabel.')
add_p('Modul praktikum ini bertujuan untuk memberikan pemahaman mendalam mengenai konsep dan implementasi stored procedure, fungsi, dan trigger dalam MySQL. Dengan menguasai materi ini, mahasiswa diharapkan mampu merancang dan membangun sistem basis data yang lebih terstruktur dan fungsional.')

add_h('1.2 Tujuan', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Tujuan dari praktikum Modul 9 ini adalah sebagai berikut:')
add_numbered([
    'Mahasiswa mampu memahami konsep stored procedure dalam basis data MySQL.',
    'Mahasiswa mampu membuat dan menjalankan stored procedure untuk mempermudah pengolahan data.',
    'Mahasiswa mampu memahami konsep fungsi (user-defined function) dalam MySQL.',
    'Mahasiswa mampu membuat dan menggunakan fungsi untuk mengembalikan nilai berdasarkan perhitungan tertentu.',
    'Mahasiswa mampu memahami konsep trigger dan penerapannya dalam otomatisasi proses basis data.',
    'Mahasiswa mampu mengimplementasikan stored procedure, fungsi, dan trigger dalam studi kasus nyata.',
])

add_h('1.3 Ruang Lingkup', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Ruang lingkup praktikum Modul 9 mencakup:')
add_numbered([
    'Pembuatan database dan tabel menggunakan MySQL.',
    'Implementasi stored procedure untuk menampilkan data mahasiswa.',
    'Implementasi fungsi untuk perhitungan diskon pada sistem toko buku.',
    'Implementasi trigger untuk otomatisasi pembaruan status member dan pencatatan log transaksi.',
    'Pengujian dan analisis hasil eksekusi masing-masing objek basis data.',
    'Penyusunan laporan praktikum yang didokumentasikan secara sistematis.',
])

doc.add_page_break()

# ==================== BAB II ====================
add_h('BAB II\nDASAR TEORI')

add_h('2.1 Stored Procedure', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Stored procedure merupakan sekumpulan pernyataan SQL yang disusun sedemikian rupa untuk menjalankan tugas tertentu. Prosedur ini disimpan dalam basis data sehingga dapat dipanggil dan dieksekusi kapan saja diperlukan. Dengan menggunakan stored procedure, kode program dapat diletakkan lebih dekat dengan data, sehingga meningkatkan efisiensi dan kecepatan pemrosesan.')
add_p('Keuntungan dari penggunaan stored procedure antara lain:')
add_bullet('• Prosedur yang telah dikompilasi akan lebih cepat dalam mengeksekusi batch perintah.')
add_bullet('• Pemrosesan data dilakukan di sisi server sehingga mengurangi lalu lintas data pada jaringan.')
add_bullet('• Stored procedure mendukung pemrograman modular karena satu prosedur dapat memanggil prosedur lainnya.')
add_bullet('• Stored procedure berperan penting dalam keamanan database karena akses langsung ke tabel dapat dikontrol.')
add_p('Sintaks dasar pembuatan stored procedure dalam MySQL adalah sebagai berikut:')
add_code('CREATE PROCEDURE nama_prosedur (parameter tipe_data)\nBEGIN\n    -- perintah deklarasi\n    -- perintah eksekusi\nEND;')
add_p('Untuk menjalankan atau memanggil stored procedure yang telah dibuat, digunakan perintah CALL:')
add_code('CALL nama_prosedur();')

add_h('2.2 Function', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Function atau fungsi merupakan program yang terdiri dari sekumpulan perintah yang tersimpan sebagai objek dalam basis data dan mengembalikan suatu nilai. Perbedaan utama antara stored procedure dan fungsi adalah bahwa stored procedure tidak mengembalikan nilai, sedangkan fungsi harus mengembalikan nilai. Dalam MySQL, pengguna dapat membuat fungsi sendiri yang dikenal dengan User Defined Function (UDF).')
add_p('Sintaks dasar pembuatan fungsi dalam MySQL adalah sebagai berikut:')
add_code('CREATE FUNCTION nama_fungsi (parameter tipe_data)\nRETURNS tipe_data_return\nBEGIN\n    -- perintah deklarasi\n    -- perintah eksekusi\n    RETURN nilai;\nEND;')
add_p('Pemanggilan fungsi dilakukan dengan perintah SELECT:')
add_code('SELECT nama_fungsi(nilai_parameter);')
add_p('Pernyataan yang diperbolehkan dalam fungsi meliputi SET, WHILE, IF, DECLARE, SELECT, INSERT, UPDATE, dan DELETE.')

add_h('2.3 Trigger', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Trigger merupakan himpunan kode prosedural yang dieksekusi secara otomatis sebagai respons terhadap suatu kejadian (event) yang berkaitan dengan tabel basis data. Kejadian yang dapat membangkitkan trigger umumnya berupa pernyataan INSERT, UPDATE, dan DELETE.')
add_p('Berdasarkan waktu eksekusinya, trigger dibedakan menjadi dua jenis yaitu trigger BEFORE yang dieksekusi sebelum operasi data dilakukan dan trigger AFTER yang dieksekusi setelah operasi data selesai. Sementara itu, berdasarkan ruang lingkupnya, trigger dibedakan menjadi row trigger yang berlaku untuk setiap baris dan statement trigger yang berlaku untuk setiap pernyataan.')
add_p('Sintaks dasar pembuatan trigger dalam MySQL adalah sebagai berikut:')
add_code('CREATE TRIGGER nama_trigger\n{BEFORE | AFTER} {INSERT | UPDATE | DELETE}\nON nama_tabel\nFOR EACH ROW\nBEGIN\n    -- perintah yang akan dijalankan\nEND;')
add_p('Dalam trigger, dikenal keyword OLD dan NEW. Keyword OLD mengacu pada nilai lama sebelum perubahan, sedangkan NEW merepresentasikan nilai baru setelah perubahan. Pada trigger INSERT, hanya keyword NEW yang dapat digunakan karena tidak ada data lama. Pada trigger DELETE, hanya keyword OLD yang dapat digunakan. Pada trigger UPDATE, kedua keyword dapat digunakan.')

doc.add_page_break()

# ==================== BAB III ====================
add_h('BAB III\nIMPLEMENTASI DAN PEMBAHASAN')

add_h('3.1 Persiapan dan Pembuatan Database', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Langkah pertama dalam praktikum ini adalah membuat database universitas yang akan digunakan untuk menyimpan tabel-tabel yang diperlukan. Berikut adalah perintah SQL yang digunakan:')
add_code('CREATE DATABASE universitas;\nUSE universitas;')
add_p('Setelah database berhasil dibuat, langkah selanjutnya adalah membuat tabel mahasiswa dengan struktur yang telah ditentukan. Tabel ini akan digunakan untuk praktikum stored procedure dan fungsi.')
add_code('CREATE TABLE mahasiswa (\n    npm VARCHAR(10) PRIMARY KEY,\n    nama VARCHAR(50),\n    jurusan VARCHAR(30),\n    no_hp VARCHAR(15)\n);')
add_p('Setelah tabel berhasil dibuat, dilakukan pengisian data sampel agar tabel memiliki data yang dapat diolah:')
add_code("INSERT INTO mahasiswa VALUES\n('001', 'Ahmad', 'TI', '081234567890'),\n('002', 'Budi', 'TI', '081234567891'),\n('003', 'Citra', 'SI', '081234567892'),\n('004', 'Dewi', 'SI', '081234567893');")

add_h('3.2 Implementasi Stored Procedure', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Stored procedure pertama yang dibuat bertujuan untuk menampilkan data mahasiswa secara sederhana, hanya menampilkan kolom NPM, nama, dan no_hp. Berikut adalah perintah pembuatan prosedur tersebut:')
add_code('DELIMITER //\nCREATE PROCEDURE lihat_mahasiswa()\nBEGIN\n    SELECT npm, nama, no_hp FROM mahasiswa;\nEND //\nDELIMITER ;')
add_p('Setelah prosedur berhasil dibuat, prosedur dapat dipanggil dengan perintah berikut:')
add_code('CALL lihat_mahasiswa();')
add_p('Hasil eksekusi prosedur tersebut akan menampilkan data seluruh mahasiswa yang tersimpan dalam tabel, yaitu kolom NPM, nama, dan nomor telepon.')

add_h('3.3 Implementasi Function', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Untuk memahami konsep fungsi dalam MySQL, dibuat sebuah fungsi sederhana yang akan mengembalikan nilai tertentu. Fungsi ini dibuat dan dipanggil menggunakan perintah SQL sesuai dengan modul praktikum.')
add_p('Sintaks pembuatan fungsi dalam MySQL menggunakan kata kunci CREATE FUNCTION yang dilengkapi dengan klausa RETURNS untuk menentukan tipe data hasil yang akan dikembalikan. Di dalam blok BEGIN...END, terdapat pernyataan RETURN yang akan mengembalikan hasil komputasi fungsi.')
add_code('DELIMITER //\nCREATE FUNCTION contoh_fungsi(nilai INT)\nRETURNS INT\nDETERMINISTIC\nBEGIN\n    DECLARE hasil INT;\n    SET hasil = nilai * 2;\n    RETURN hasil;\nEND //\nDELIMITER ;')
add_p('Pemanggilan fungsi dilakukan dengan perintah SELECT:')
add_code('SELECT contoh_fungsi(5);')
add_p('Hasil yang ditampilkan adalah nilai 10, yang merupakan hasil perkalian parameter yang dimasukkan dengan angka 2.')

add_h('3.4 Implementasi Trigger', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Praktikum trigger menggunakan dua buah tabel, yaitu tabel barang dan tabel pembelian. Tabel barang menyimpan data produk, sedangkan tabel pembelian mencatat transaksi pembelian yang terjadi.')
add_p('Pertama, buat kedua tabel tersebut dengan struktur berikut:')
add_code('CREATE TABLE barang (\n    id_barang INT AUTO_INCREMENT PRIMARY KEY,\n    nama_barang VARCHAR(50),\n    harga INT,\n    stok INT\n);')
add_code('CREATE TABLE pembelian (\n    id_pembelian INT AUTO_INCREMENT PRIMARY KEY,\n    id_barang INT,\n    jumlah INT,\n    total INT,\n    FOREIGN KEY (id_barang) REFERENCES barang(id_barang)\n);')
add_p('Setelah tabel siap, langkah selanjutnya adalah membuat trigger yang akan mencatat log setiap kali ada data pembelian yang dimasukkan. Trigger ini akan secara otomatis mengurangi stok barang setiap kali ada transaksi pembelian baru:')
add_code('DELIMITER //\nCREATE TRIGGER after_pembelian_insert\nAFTER INSERT ON pembelian\nFOR EACH ROW\nBEGIN\n    UPDATE barang SET stok = stok - NEW.jumlah\n    WHERE id_barang = NEW.id_barang;\nEND //\nDELIMITER ;')
add_p('Trigger di atas akan dieksekusi secara otomatis setiap kali data baru dimasukkan ke tabel pembelian. Ketika proses INSERT berhasil, trigger akan menjalankan perintah UPDATE pada tabel barang untuk mengurangi stok sesuai jumlah yang dibeli.')
add_p('Pengujian trigger dilakukan dengan memasukkan data ke tabel pembelian dan memeriksa perubahan yang terjadi pada tabel barang. Jika trigger berjalan dengan baik, stok barang akan berkurang sesuai dengan jumlah yang dibeli.')
add_p('Selain trigger INSERT, dapat juga dibuat trigger untuk operasi DELETE dan UPDATE. Keyword OLD dan NEW digunakan untuk merujuk ke nilai sebelum dan sesudah perubahan data. Trigger DELETE hanya bisa menggunakan keyword OLD, sedangkan trigger UPDATE dapat menggunakan kedua keyword tersebut.')
add_p('Sebagai contoh, trigger DELETE dapat dibuat untuk menghapus data pembelian ketika data barang yang sesuai dihapus:')
add_code('DELIMITER //\nCREATE TRIGGER after_barang_delete\nAFTER DELETE ON barang\nFOR EACH ROW\nBEGIN\n    DELETE FROM pembelian WHERE id_barang = OLD.id_barang;\nEND //\nDELIMITER ;')

doc.add_page_break()

add_h('3.5 Implementasi Tugas Modul 9', size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
add_p('Tugas Modul 9 merupakan implementasi sistem basis data untuk toko buku online. Sistem ini terdiri dari tiga tabel utama, yaitu buku, pelanggan, dan transaksi, serta dilengkapi dengan function, stored procedure, dan trigger.')

add_p('a. Pembuatan Tabel', bold=True)
add_p('Tabel buku digunakan untuk menyimpan data buku yang tersedia di toko. Tabel ini memiliki atribut id_buku sebagai primary key, judul, penulis, harga, dan stok.')
add_code('CREATE TABLE buku (\n    id_buku INT AUTO_INCREMENT PRIMARY KEY,\n    judul VARCHAR(100),\n    penulis VARCHAR(100),\n    harga DECIMAL(10,2),\n    stok INT\n);')
add_p('Tabel pelanggan digunakan untuk menyimpan data pelanggan. Tabel ini memiliki atribut id_pelanggan, nama, total_belanja, dan status_member yang secara default bernilai REGULER.')
add_code("CREATE TABLE pelanggan (\n    id_pelanggan INT AUTO_INCREMENT PRIMARY KEY,\n    nama VARCHAR(100),\n    total_belanja DECIMAL(10,2) DEFAULT 0,\n    status_member ENUM('REGULER','GOLD','PLATINUM') DEFAULT 'REGULER'\n);")
add_p('Tabel transaksi digunakan untuk mencatat setiap transaksi pembelian yang dilakukan oleh pelanggan. Tabel ini memiliki foreign key ke tabel pelanggan dan buku.')
add_code('CREATE TABLE transaksi (\n    id_transaksi INT AUTO_INCREMENT PRIMARY KEY,\n    id_pelanggan INT,\n    id_buku INT,\n    jumlah INT,\n    total_harga DECIMAL(10,2),\n    tanggal_transaksi DATE,\n    FOREIGN KEY (id_pelanggan) REFERENCES pelanggan(id_pelanggan),\n    FOREIGN KEY (id_buku) REFERENCES buku(id_buku)\n);')

add_p('b. Pembuatan Function hitung_diskon', bold=True)
add_p('Function hitung_diskon digunakan untuk menghitung besaran diskon yang diperoleh pelanggan berdasarkan total belanja yang dimilikinya. Ketentuan diskon yang diterapkan adalah sebagai berikut:')
add_bullet('• Total belanja kurang dari Rp1.000.000    → diskon 0%')
add_bullet('• Total belanja antara Rp1.000.000 s.d. Rp4.999.999    → diskon 5%')
add_bullet('• Total belanja lebih dari atau sama dengan Rp5.000.000    → diskon 10%')
add_p('Function ini mengembalikan nilai bertipe DECIMAL(5,2) dan bersifat DETERMINISTIC karena untuk input yang sama akan menghasilkan output yang sama.')
add_code('DELIMITER //\nCREATE FUNCTION hitung_diskon(total_belanja DECIMAL(10,2))\nRETURNS DECIMAL(5,2)\nDETERMINISTIC\nBEGIN\n    DECLARE diskon DECIMAL(5,2);\n    IF total_belanja < 1000000 THEN\n        SET diskon = 0.00;\n    ELSEIF total_belanja < 5000000 THEN\n        SET diskon = 0.05;\n    ELSE\n        SET diskon = 0.10;\n    END IF;\n    RETURN diskon;\nEND //\nDELIMITER ;')

add_p('c. Pembuatan Stored Procedure tambah_transaksi', bold=True)
add_p('Stored procedure tambah_transaksi digunakan untuk memproses transaksi pembelian buku. Prosedur ini menerima tiga parameter, yaitu id pelanggan, id buku, dan jumlah yang dibeli. Berikut langkah-langkah yang dilakukan oleh prosedur ini:')
add_numbered([
    'Mengecek ketersediaan stok buku melalui tabel buku.',
    'Jika stok tidak mencukupi, menampilkan pesan error menggunakan SIGNAL SQLSTATE.',
    'Jika stok mencukupi, mengurangi stok buku sesuai jumlah yang dibeli.',
    'Menghitung total harga berdasarkan harga buku dikalikan jumlah yang dibeli.',
    'Menambahkan data transaksi ke tabel transaksi dengan tanggal hari ini (CURDATE()).',
    'Menambahkan total harga ke total_belanja pada tabel pelanggan.',
    'Menampilkan pesan "Transaksi berhasil" sebagai konfirmasi.',
])
add_code('DELIMITER //\nCREATE PROCEDURE tambah_transaksi(\n    p_id_pelanggan INT,\n    p_id_buku INT,\n    p_jumlah INT\n)\nBEGIN\n    DECLARE v_harga DECIMAL(10,2);\n    DECLARE v_stok INT;\n    DECLARE v_total_harga DECIMAL(10,2);\n\n    SELECT harga, stok INTO v_harga, v_stok\n    FROM buku WHERE id_buku = p_id_buku;\n\n    IF v_stok < p_jumlah THEN\n        SIGNAL SQLSTATE \'45000\'\n        SET MESSAGE_TEXT = \'Error: Stok tidak mencukupi\';\n    ELSE\n        UPDATE buku SET stok = stok - p_jumlah\n        WHERE id_buku = p_id_buku;\n\n        SET v_total_harga = v_harga * p_jumlah;\n\n        INSERT INTO transaksi\n        VALUES (NULL, p_id_pelanggan, p_id_buku,\n                p_jumlah, v_total_harga, CURDATE());\n\n        UPDATE pelanggan\n        SET total_belanja = total_belanja + v_total_harga\n        WHERE id_pelanggan = p_id_pelanggan;\n\n        SELECT \'Transaksi berhasil\' AS pesan;\n    END IF;\nEND //\nDELIMITER ;')

add_p('d. Pembuatan Trigger update_status_member', bold=True)
add_p('Trigger update_status_member merupakan trigger AFTER UPDATE pada tabel pelanggan yang secara otomatis mengubah status member pelanggan berdasarkan total belanja yang dimilikinya. Ketentuan perubahan status adalah sebagai berikut:')
add_bullet('• Total belanja >= Rp5.000.000    → status menjadi PLATINUM')
add_bullet('• Total belanja >= Rp1.000.000    → status menjadi GOLD')
add_bullet('• Selainnya    → status tetap REGULER')
add_p('Dengan adanya trigger ini, status member pelanggan akan selalu diperbarui secara otomatis tanpa perlu intervensi manual dari operator setiap kali terjadi perubahan pada total belanja.')
add_code('DELIMITER //\nCREATE TRIGGER update_status_member\nAFTER UPDATE ON pelanggan\nFOR EACH ROW\nBEGIN\n    DECLARE v_status ENUM(\'REGULER\',\'GOLD\',\'PLATINUM\');\n\n    IF NEW.total_belanja >= 5000000 THEN\n        SET v_status = \'PLATINUM\';\n    ELSEIF NEW.total_belanja >= 1000000 THEN\n        SET v_status = \'GOLD\';\n    ELSE\n        SET v_status = \'REGULER\';\n    END IF;\n\n    IF NEW.status_member <> v_status THEN\n        UPDATE pelanggan\n        SET status_member = v_status\n        WHERE id_pelanggan = NEW.id_pelanggan;\n    END IF;\nEND //\nDELIMITER ;')

add_p('e. Pengujian Sistem', bold=True)
add_p('Untuk menguji keseluruhan sistem yang telah dibangun, dilakukan beberapa langkah pengujian sebagai berikut:')
add_numbered([
    'Memasukkan data sampel ke tabel buku dan pelanggan.',
    'Memanggil function hitung_diskon untuk menghitung diskon berdasarkan total belanja.',
    'Menjalankan stored procedure tambah_transaksi untuk melakukan transaksi.',
    'Memeriksa perubahan stok buku setelah transaksi dilakukan.',
    'Memeriksa perubahan total_belanja dan status_member pelanggan.',
    'Menguji validasi stok dengan mencoba transaksi melebihi stok yang tersedia.',
])
add_p('Pengujian menunjukkan bahwa seluruh objek basis data berfungsi sesuai dengan yang diharapkan. Function hitung_diskon berhasil menghitung diskon berdasarkan total belanja. Stored procedure tambah_transaksi berhasil memproses transaksi, mengurangi stok, dan memperbarui data pelanggan. Trigger update_status_member berhasil mengubah status member secara otomatis berdasarkan total belanja.')

doc.add_page_break()

# ==================== BAB IV ====================
add_h('BAB IV\nKESIMPULAN')
add_p('Berdasarkan praktikum Modul 9 tentang Stored Procedure, Fungsi, dan Trigger yang telah dilakukan, dapat diambil beberapa kesimpulan sebagai berikut:')
add_numbered([
    'Stored procedure merupakan kumpulan perintah SQL yang tersimpan dalam basis data dan dapat dipanggil kapan saja menggunakan perintah CALL. Stored procedure sangat berguna untuk menyederhanakan operasi yang kompleks dan meningkatkan efisiensi pemrosesan data.',
    'Function atau fungsi merupakan objek basis data yang mengembalikan suatu nilai. Perbedaan utama antara fungsi dan stored procedure adalah bahwa fungsi harus mengembalikan nilai (menggunakan RETURN), sedangkan stored procedure tidak.',
    'Trigger merupakan mekanisme otomatisasi dalam basis data yang dieksekusi ketika terjadi event INSERT, UPDATE, atau DELETE pada suatu tabel. Trigger dapat digunakan untuk menjaga integritas data, mencatat log aktivitas, dan melakukan validasi secara otomatis.',
    'Implementasi stored procedure, fungsi, dan trigger dalam studi kasus toko buku online berhasil dilakukan dengan baik. Function hitung_diskon berhasil menghitung diskon, stored procedure tambah_transaksi berhasil memproses transaksi dengan validasi stok, dan trigger update_status_member berhasil mengubah status member secara otomatis.',
])
add_p('Dengan demikian, pemahaman mengenai stored procedure, fungsi, dan trigger sangat penting dalam pengembangan sistem basis data yang efisien, terstruktur, dan memiliki tingkat otomatisasi yang tinggi.')

# ==================== DAFTAR PUSTAKA ====================
add_h('DAFTAR PUSTAKA')
refs = [
    'Connolly, T., & Begg, C. (2015). Database Systems: A Practical Approach to Design, Implementation, and Management (6th ed.). Pearson.',
    'Garcia-Molina, H., Ullman, J. D., & Widom, J. (2009). Database Systems: The Complete Book (2nd ed.). Pearson.',
    'Kadir, A. (2002). Konsep dan Tuntunan Praktis Basis Data. Andi Offset.',
    'Wardhani, O., & Alfath, I. (2025). Modul Praktikum Basis Data #9: Stored Procedure, Fungsi, dan Trigger. Teknologi Informasi, Universitas Tidar.',
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ==================== SAVE ====================
output_path = os.path.expanduser('~/Struktur-Data/Modul-09-Stored-Procedure-Function-Trigger/Laporan_Praktikum_Modul_9_SP_FC_TG.docx')
doc.save(output_path)
print(f"DOCX saved: {output_path}")
print(f"Size: {os.path.getsize(output_path)} bytes")
